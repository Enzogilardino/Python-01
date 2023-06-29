from docx import Document
import tkinter as tk

data_pesanan = []

def tambah_pesanan():
    nama = entry_nama.get()
    harga = entry_harga.get()
    pesanan = entry_pesanan.get()
    data_pesanan.append((nama, harga, pesanan))
    entry_nama.delete(0, tk.END)
    entry_harga.delete(0, tk.END)
    entry_pesanan.delete(0, tk.END)

def tulis_ke_word():
    document = Document()

    document.add_heading('Invoice', 0)

    for i, data in enumerate(data_pesanan, start=1):
        nama, harga, pesanan = data
        document.add_paragraph(f'\nData untuk pesanan ke-{i}:')
        document.add_paragraph('Nama: ' + nama)
        document.add_paragraph('Harga barang: ' + harga)
        document.add_paragraph('Pesanan: ' + pesanan)

    document.save('invoice.docx')
    label_status.config(text="Data berhasil ditulis ke file invoice.docx")

# Membuat jendela Tkinter
window = tk.Tk()
window.title("Input Data")
window.geometry("300x300")

# Membuat label dan entry untuk setiap data
label_nama = tk.Label(window, text="Nama:")
label_nama.pack()
entry_nama = tk.Entry(window)
entry_nama.pack()

label_harga = tk.Label(window, text="Harga barang:")
label_harga.pack()
entry_harga = tk.Entry(window)
entry_harga.pack()

label_pesanan = tk.Label(window, text="Pesanan:")
label_pesanan.pack()
entry_pesanan = tk.Entry(window)
entry_pesanan.pack()

# Tombol untuk menambah pesanan
button_tambah = tk.Button(window, text="Tambah Pesanan", command=tambah_pesanan)
button_tambah.pack()

# Tombol untuk menulis ke Word
button_tulis = tk.Button(window, text="Tulis ke Word", command=tulis_ke_word)
button_tulis.pack()

# Label untuk menampilkan status
label_status = tk.Label(window, text="")
label_status.pack()

window.mainloop()
