from docx import Document
import tkinter as tk
from tkinter import filedialog

def input_data():
    data_pesanan = []
    jumlah_pesanan = int(entry_jumlah.get())

    for i in range(jumlah_pesanan):
        window.withdraw()  # Menyembunyikan jendela utama sementara
        child_window = tk.Toplevel(window)
        child_window.title(f"Data Pesanan {i+1}")
        child_window.geometry("300x200")

        # Membuat label dan entry untuk setiap data
        label_nama = tk.Label(child_window, text="Nama:")
        label_nama.pack()
        entry_nama = tk.Entry(child_window)
        entry_nama.pack()

        label_harga = tk.Label(child_window, text="Harga barang:")
        label_harga.pack()
        entry_harga = tk.Entry(child_window)
        entry_harga.pack()

        label_pesanan = tk.Label(child_window, text="Pesanan:")
        label_pesanan.pack()
        entry_pesanan = tk.Entry(child_window)
        entry_pesanan.pack()

        # Menyimpan data pesanan
        button_simpan = tk.Button(child_window, text="Simpan", command=lambda: simpan_data(entry_nama.get(), entry_harga.get(), entry_pesanan.get()))
        button_simpan.pack()

        def simpan_data(nama, harga, pesanan):
            data_pesanan.append((nama, harga, pesanan))
            child_window.destroy()
            if len(data_pesanan) == jumlah_pesanan:
                window.deiconify()  # Menampilkan kembali jendela utama setelah semua data tersimpan

    def tulis_ke_word():
        file_path = filedialog.askopenfilename(initialdir="/", title="Pilih File", filetypes=(("Semua file", "*.*"),))

        if file_path:
            document = Document()
            
            document.add_heading('Invoice', 0)
            
            for i, data in enumerate(data_pesanan, start=1):
                nama, harga, pesanan = data
                document.add_paragraph(f'\nData untuk pesanan ke-{i}:')
                document.add_paragraph('Nama: ' + nama)
                document.add_paragraph('Harga barang: ' + harga)
                document.add_paragraph('Pesanan: ' + pesanan)
            
            document.save(file_path)
            label_status.config(text="Data berhasil ditulis ke file " + file_path)
        else:
            label_status.config(text="Tidak ada file yang dipilih")

    # Membuat jendela Tkinter
    window = tk.Tk()
    window.title("Input Data")
    window.geometry("300x250")

    # Membuat label dan entry untuk jumlah pesanan
    label_jumlah = tk.Label(window, text="Jumlah Pesanan:")
    label_jumlah.pack()
    entry_jumlah = tk.Entry(window)
    entry_jumlah.pack()

    # Tombol untuk memilih file dan menulis ke Word
    button_file = tk.Button(window, text="Pilih File", command=tulis_ke_word)
    button_file.pack()

    # Label untuk menampilkan status
    label_status = tk.Label(window, text="")
    label_status.pack()

    window.mainloop()

input_data()
